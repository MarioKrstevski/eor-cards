"""Chunker service — parse .docx study notes and chunk them semantically using Claude API."""

import base64
import io
import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

import anthropic
from docx import Document
from lxml import etree
from PIL import Image

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"

INDENT_LEVEL_BASE = 3

IMG_DATA_URI_RE = re.compile(r'<img[^>]+src="(data:image/[^"]+)"', re.IGNORECASE)


def extract_images_from_paragraph(para, doc, img_dir, doc_prefix, para_idx):
    """Extract images from a paragraph, save to disk, return list of image info dicts."""
    images = []
    drawings = para._element.findall(f".//{{{WP_NS}}}inline") + para._element.findall(
        f".//{{{WP_NS}}}anchor"
    )
    for img_idx, drawing in enumerate(drawings):
        blip = drawing.find(f".//{{{A_NS}}}blip")
        if blip is None:
            continue
        embed_id = blip.get(f"{{{R_NS}}}embed")
        if embed_id is None:
            continue
        try:
            rel = doc.part.rels[embed_id]
            img_blob = rel.target_part.blob
            ext = os.path.splitext(rel.target_part.partname)[1] or ".png"
            filename = f"{doc_prefix}_p{para_idx}_i{img_idx}{ext}"
            filepath = os.path.join(img_dir, filename)
            with open(filepath, "wb") as f:
                f.write(img_blob)
            img = Image.open(io.BytesIO(img_blob))
            if img.width > 800:
                ratio = 800 / img.width
                img = img.resize((800, int(img.height * ratio)), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            fmt = "PNG" if ext.lower() in (".png", ".gif") else "JPEG"
            img.save(buf, format=fmt)
            b64 = base64.b64encode(buf.getvalue()).decode()
            media_type = "image/png" if fmt == "PNG" else "image/jpeg"
            data_uri = f"data:{media_type};base64,{b64}"
            images.append({
                "filename": filename,
                "filepath": filepath,
                "data_uri": data_uri,
                "base64": b64,
                "media_type": media_type,
                "para_index": para_idx,
            })
        except Exception as e:
            logger.warning("Could not extract image %d from para %d: %s", img_idx, para_idx, e)
    return images


def get_paragraph_level(para):
    """Get the nesting level of a paragraph. Returns (level, source)."""
    numPr = para._element.find(f".//{{{WML_NS}}}numPr")
    if numPr is not None:
        ilvl_elem = numPr.find(f"{{{WML_NS}}}ilvl")
        if ilvl_elem is not None:
            val = ilvl_elem.get(f"{{{WML_NS}}}val")
            if val is not None:
                return int(val), "ilvl"
    left_indent = para.paragraph_format.left_indent
    if left_indent and left_indent > 0:
        return left_indent, "indent_raw"
    return 0, "none"


def extract_runs_with_formatting(para):
    """Extract text runs with bold/italic information."""
    runs = []
    for run in para.runs:
        if not run.text:
            continue
        runs.append({"text": run.text, "bold": bool(run.bold), "italic": bool(run.italic)})
    return runs


def runs_to_html(runs):
    """Convert formatted runs to HTML string."""
    html = ""
    for r in runs:
        text = r["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if r["bold"] and r["italic"]:
            html += f"<b><i>{text}</i></b>"
        elif r["bold"]:
            html += f"<b>{text}</b>"
        elif r["italic"]:
            html += f"<i>{text}</i>"
        else:
            html += text
    return html


def parse_docx(docx_path: str, img_dir: str):
    """Parse a .docx file into structured elements."""
    doc = Document(docx_path)
    os.makedirs(img_dir, exist_ok=True)
    doc_prefix = Path(docx_path).stem.replace(" ", "_")[:40]
    elements = []
    all_images = []

    indent_values = set()
    for para in doc.paragraphs:
        level, source = get_paragraph_level(para)
        if source == "indent_raw":
            indent_values.add(level)

    sorted_indents = sorted(indent_values)
    indent_to_level = {indent_val: INDENT_LEVEL_BASE + i for i, indent_val in enumerate(sorted_indents)}

    body = doc.element.body
    para_idx = 0
    table_idx = 0
    body_order = []
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p":
            body_order.append(("para", para_idx))
            para_idx += 1
        elif tag == "tbl":
            body_order.append(("table", table_idx))
            table_idx += 1

    para_elements = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        runs = extract_runs_with_formatting(para)
        level, source = get_paragraph_level(para)

        if source == "indent_raw":
            level = indent_to_level[level]

        is_section_bullet = text.startswith("§")
        if is_section_bullet:
            text = text.lstrip("§").strip()
            if runs and runs[0]["text"].startswith("§"):
                runs[0]["text"] = runs[0]["text"].lstrip("§").strip()
                if not runs[0]["text"]:
                    runs = runs[1:]

        bold_terms = [r["text"].strip() for r in runs if r["bold"] and r["text"].strip()]
        all_text_bold = all(r["bold"] for r in runs if r["text"].strip()) if runs else False
        is_heading = source == "ilvl" and level <= 2 and text and all_text_bold

        images = extract_images_from_paragraph(para, doc, img_dir, doc_prefix, i)
        all_images.extend(images)

        html = runs_to_html(runs)

        para_elements[i] = {
            "type": "heading" if is_heading else "bullet" if (source in ("ilvl", "indent_raw") and level >= 3) or is_section_bullet else "paragraph",
            "text": text,
            "html": html,
            "bold_terms": bold_terms,
            "level": level,
            "para_index": i,
            "images": [{"data_uri": img["data_uri"], "media_type": img["media_type"]} for img in images],
            "is_empty": not text and not images,
        }

    table_data = {}
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_runs = []
                for p in cell.paragraphs:
                    cell_runs.extend(extract_runs_with_formatting(p))
                cells.append({"text": cell.text.strip(), "html": runs_to_html(cell_runs)})
            rows.append(cells)
        table_data[t_idx] = {
            "type": "table",
            "rows": rows,
            "text": "\n".join(" | ".join(cell["text"] for cell in row) for row in rows),
            "html": table_to_html(rows),
            "bold_terms": [],
            "level": 0,
            "para_index": -1,
            "images": [],
            "is_empty": False,
        }

    for item_type, idx in body_order:
        if item_type == "para":
            elem = para_elements.get(idx)
            if elem and not elem["is_empty"]:
                elements.append(elem)
        elif item_type == "table":
            tbl = table_data.get(idx)
            if tbl:
                tbl["para_index"] = len(elements)
                elements.append(tbl)

    return elements, all_images


def table_to_html(rows):
    """Convert table rows to HTML."""
    html = '<table border="1" cellpadding="4" cellspacing="0">'
    for r_idx, row in enumerate(rows):
        html += "<tr>"
        tag = "th" if r_idx == 0 else "td"
        for cell in row:
            html += f"<{tag}>{cell['html']}</{tag}>"
        html += "</tr>"
    html += "</table>"
    return html


def elements_to_document_html(elements):
    """Generate full document HTML from elements (for the viewer left panel)."""
    html_parts = []
    for elem in elements:
        if elem["type"] == "heading":
            level = min(elem.get("level", 2), 4)
            tag = f"h{level + 1}"
            html_parts.append(f"<{tag}>{elem['html']}</{tag}>")
        elif elem["type"] == "table":
            html_parts.append(elem["html"])
        elif elem["type"] == "bullet":
            html_parts.append(f'<li class="level-{elem.get("level", 3)}">{elem["html"]}</li>')
        else:
            html_parts.append(f"<p>{elem['html']}</p>")
        for img in elem.get("images", []):
            html_parts.append(f'<img src="{img["data_uri"]}" class="doc-image" />')
    return "\n".join(html_parts)


def build_claude_chunking_prompt(elements, images, rules_md_path=None):
    """Build the prompt for Claude to decide chunk boundaries."""
    elem_summaries = []
    for i, elem in enumerate(elements):
        summary = {
            "index": i, "type": elem["type"], "level": elem.get("level", 0),
            "text": elem["text"][:500], "bold_terms": elem.get("bold_terms", []),
            "has_images": bool(elem.get("images")),
        }
        if elem["type"] == "table":
            summary["text"] = elem["text"][:300]
        elem_summaries.append(summary)

    rule_info = ""
    if rules_md_path and os.path.exists(rules_md_path):
        rule_info = """
The following rule subsets exist for later card generation. Tag each chunk with applicable subsets:
- "mechanism_splitting": For paragraphs explaining cause-effect chains
- "symptom_clusters": For bullet lists of symptoms from the same mechanism
- "clinical_decision_tree": For branching decision logic (if X then Y, else Z)
- "treatment_contrast": For comparing treatments or medications
- "cloze_boundaries": Applies to all chunks (always include)
- "timing_markers": For content with specific timing/duration info
- "cause_pathology_lists": For lists of causes or underlying pathologies
- "exam_findings": For physical exam or diagnostic findings
"""

    prompt = f"""You are analyzing a medical study document that has been parsed into structured elements.
Your job is to group these elements into semantic "chunks" — each chunk is a self-contained study unit
that a student could review independently.

RULES FOR CHUNKING:
1. Each chunk must be semantically coherent — it covers one concept or closely related set of concepts
2. A heading with its sub-bullets/content stays together as one chunk
3. A short introductory paragraph followed by supporting bullets should be ONE chunk (not split)
4. Tables are their own chunk unless they're small and directly support adjacent text
5. Don't make chunks too large (max ~15-20 elements) or too small (min 2-3 elements unless it's a heading-only section)
6. Bold text marks "things to remember" — don't split bold content from its context
7. If a paragraph or bullet introduces sub-items, keep them together
8. IMPORTANT: Top-level category/title headings (e.g., "CARDIOVASCULAR", "Atrial fibrillation" at the very start) are document context, NOT study content. Do NOT create a chunk for them. Skip them entirely — start chunking from the first actual study content.

CONTENT TYPE CLASSIFICATION — assign one per chunk:
- "paragraph": Primarily flowing text
- "bullet-list": Primarily bullets/sub-bullets
- "mixed-paragraph-bullets": Intro paragraph + supporting bullets
- "table": Contains a table
- "chart-image": Contains a chart, EKG, or data visualization
- "text-screenshot-image": Contains a screenshot of text from another source
- "reference-image": Contains a diagram or illustration

{rule_info}

Here are the document elements (index, type, level, text preview, bold terms):

{json.dumps(elem_summaries, indent=2)}

Return a JSON array of chunks. Each chunk:
{{
  "chunk_index": <sequential 0-based>,
  "element_range": [<start_element_index>, <end_element_index_inclusive>],
  "heading": "<descriptive heading for this chunk>",
  "content_type": "<one of the types above>",
  "rule_subset": ["cloze_boundaries", "<other applicable rules>"],
  "rationale": "<brief explanation of why these elements belong together>"
}}

Return ONLY the JSON array, no other text."""
    return prompt


def call_claude_for_chunking(elements: list, images: list, rules_md_path: Optional[str], client: anthropic.Anthropic, model: str = "claude-sonnet-4-6") -> tuple[list, dict]:
    """Call Claude API to determine chunk boundaries. Client is injected (not created here).
    Returns (chunks_list, usage_dict) where usage_dict = {"input_tokens": ..., "output_tokens": ...}.
    """
    prompt = build_claude_chunking_prompt(elements, images, rules_md_path)
    content = [{"type": "text", "text": prompt}]
    for img in images[:10]:
        if not img.get("base64"):
            continue
        content.append({"type": "image", "source": {"type": "base64", "media_type": img["media_type"], "data": img["base64"]}})
        content.append({"type": "text", "text": f"[Image from paragraph {img['para_index']}]"})
    response = client.messages.create(
        model=model,
        max_tokens=16384,
        temperature=0,
        messages=[{"role": "user", "content": content}],
    )
    usage = {
        "input_tokens": response.usage.input_tokens,
        "output_tokens": response.usage.output_tokens,
    }
    if response.stop_reason == "max_tokens":
        raise ValueError("Chunking response was truncated — document may be too large for single-pass chunking")
    response_text = response.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = re.sub(r"^```\w*\n?", "", response_text)
        response_text = re.sub(r"\n?```$", "", response_text)
    try:
        return json.loads(response_text), usage
    except json.JSONDecodeError as e:
        raise ValueError(f"Claude returned non-JSON chunking response: {e}\nResponse: {response_text[:200]}") from e


def assemble_chunks(elements: list, claude_chunks: list) -> list:
    """Assemble final chunk objects from elements and Claude's chunk boundaries."""
    result_chunks = []
    for cc in claude_chunks:
        start, end = cc["element_range"]
        chunk_elements = elements[start: end + 1]
        source_text_parts = []
        source_html_parts = []
        all_bold_terms = []
        all_images = []
        for elem in chunk_elements:
            source_text_parts.append(elem["text"])
            if elem["type"] == "heading":
                level = min(elem.get("level", 2), 4)
                source_html_parts.append(f"<h{level + 1}>{elem['html']}</h{level + 1}>")
            elif elem["type"] == "table":
                source_html_parts.append(elem["html"])
            elif elem["type"] == "bullet":
                source_html_parts.append(
                    f'<div class="bullet level-{elem.get("level", 3)}">{elem["html"]}</div>'
                )
            else:
                source_html_parts.append(f"<p>{elem['html']}</p>")
            all_bold_terms.extend(elem.get("bold_terms", []))
            for img in elem.get("images", []):
                all_images.append({
                    "data_uri": img["data_uri"],
                    "type": cc.get("content_type", "reference-image")
                    if "image" in cc.get("content_type", "")
                    else "reference-image",
                    "ocr_text": None,
                    "position": f"in_element_{elem['para_index']}",
                })
                source_html_parts.append(f'<img src="{img["data_uri"]}" class="doc-image" />')
        seen = set()
        unique_bold = []
        for t in all_bold_terms:
            if t not in seen:
                seen.add(t)
                unique_bold.append(t)
        source_html = "\n".join(source_html_parts)
        img_match = IMG_DATA_URI_RE.search(source_html)
        result_chunks.append({
            "chunk_index": cc["chunk_index"],
            "heading": cc["heading"],
            "content_type": cc["content_type"],
            "rule_subset": cc.get("rule_subset", ["cloze_boundaries"]),
            "source_text": "\n".join(source_text_parts),
            "source_html": source_html,
            "bold_terms": unique_bold,
            "element_range": cc["element_range"],
            "images": all_images,
            "ref_img": img_match.group(1) if img_match else None,
        })
    return result_chunks


def _parse_css_class_margins(html: str) -> dict:
    """Extract left-margin values (px) keyed by CSS selector from the <style> block.

    Pages generates rules like:
        li.li1 {margin: 0.0px 0.0px 0.0px 0.0px; ...}
        li.li2 {margin: 0.0px 0.0px 0.0px 28.0px; ...}
    The fourth value of the shorthand margin is the left margin.
    """
    result: dict = {}
    style_block = re.search(r'<style[^>]*>(.*?)</style>', html, re.DOTALL | re.I)
    if not style_block:
        return result
    css = style_block.group(1)
    for rule in re.finditer(r'([\w.#-]+)\s*\{([^}]+)\}', css):
        selector = rule.group(1).strip()
        props = rule.group(2)
        # margin shorthand: top right bottom left
        m = re.search(
            r'margin\s*:\s*([\d.]+)(?:px|pt)\s+([\d.]+)(?:px|pt)\s+([\d.]+)(?:px|pt)\s+([\d.]+)(px|pt)',
            props, re.I
        )
        if m:
            value = float(m.group(4))
            unit = m.group(5).lower()
            if unit == 'pt':
                value = value * 1.333  # pt → px
            result[selector] = value
            continue
        # explicit margin-left / padding-left
        ml = re.search(r'(?:margin|padding)-left\s*:\s*([\d.]+)(px|pt)', props, re.I)
        if ml:
            value = float(ml.group(1))
            if ml.group(2).lower() == 'pt':
                value = value * 1.333
            result[selector] = value
    return result


def parse_html_to_elements(html: str) -> tuple[list, list]:
    """Parse clipboard HTML (Pages, Word, Google Docs, etc.) into the same element format as parse_docx."""
    import html as html_lib
    from bs4 import BeautifulSoup, Tag, NavigableString

    # Extract just the clipboard fragment if Word/Office markers are present
    frag_match = re.search(r'<!--StartFragment-->(.*?)<!--EndFragment-->', html, re.DOTALL)
    if frag_match:
        html = frag_match.group(1)

    # Strip Word conditional comments: <!--[if ...]>...<![endif]--> and <![if ...]>...<![endif]>
    # These wrap bullet markers (○, ■, etc.) and leak as plain text through html.parser
    html = re.sub(r'<!--\[if[^\]]*\]>.*?<!\[endif\]-->', '', html, flags=re.DOTALL)
    html = re.sub(r'<!\[if[^\]]*\]>.*?<!\[endif\]>', '', html, flags=re.DOTALL)

    # Pre-parse CSS class → left-margin mapping (used for Pages indent detection)
    css_margins = _parse_css_class_margins(html)

    soup = BeautifulSoup(html, 'html.parser')
    elements: list = []
    all_images: list = []
    para_idx = 0

    def node_to_clean_html(node) -> str:
        """Recursively convert a BS4 node to HTML keeping only b/i formatting."""
        if isinstance(node, NavigableString):
            return html_lib.escape(str(node))
        tag = node.name or ''
        # Skip XML/Office namespace tags (o:p, w:*, etc.)
        if ':' in tag:
            return ''
        inner = ''.join(node_to_clean_html(c) for c in node.children)
        if tag in ('b', 'strong'):
            return f'<b>{inner}</b>' if inner.strip() else inner
        if tag in ('i', 'em'):
            return f'<i>{inner}</i>' if inner.strip() else inner
        # Honour inline bold/italic via style attribute
        style = (node.get('style', '') if hasattr(node, 'get') else '').replace(' ', '').lower()
        if 'font-weight:bold' in style or 'font-weight:700' in style:
            return f'<b>{inner}</b>' if inner.strip() else inner
        if 'font-style:italic' in style:
            return f'<i>{inner}</i>' if inner.strip() else inner
        return inner

    def extract_bold_terms(node) -> list:
        terms = []
        for b in node.find_all(['b', 'strong']):
            t = b.get_text().strip()
            if t:
                terms.append(t)
        for span in node.find_all('span'):
            style = (span.get('style', '') or '').replace(' ', '').lower()
            if 'font-weight:bold' in style or 'font-weight:700' in style:
                t = span.get_text().strip()
                if t:
                    terms.append(t)
        return terms

    def extract_images_from_node(node) -> list:
        imgs = []
        for img_tag in node.find_all('img'):
            src = img_tag.get('src', '')
            if not src.startswith('data:'):
                continue
            m = re.match(r'data:(image/[^;]+);base64,(.+)', src, re.DOTALL)
            if not m:
                continue
            media_type = m.group(1)
            b64_data = m.group(2).strip()
            try:
                img_bytes = base64.b64decode(b64_data)
                # Detect real format from magic bytes — source may mislabel (e.g. GIF as PNG)
                if img_bytes[:6] in (b'GIF87a', b'GIF89a'):
                    media_type = 'image/gif'
                elif img_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                    media_type = 'image/png'
                elif img_bytes[:2] == b'\xff\xd8':
                    media_type = 'image/jpeg'
                elif img_bytes[:4] == b'RIFF' and img_bytes[8:12] == b'WEBP':
                    media_type = 'image/webp'
                pil_img = Image.open(io.BytesIO(img_bytes))
                if pil_img.width > 800:
                    ratio = 800 / pil_img.width
                    pil_img = pil_img.resize((800, int(pil_img.height * ratio)), Image.Resampling.LANCZOS)
                    buf = io.BytesIO()
                    fmt = "PNG" if media_type == "image/png" else "GIF" if media_type == "image/gif" else "JPEG"
                    pil_img.save(buf, format=fmt)
                    b64_data = base64.b64encode(buf.getvalue()).decode()
                data_uri = f"data:{media_type};base64,{b64_data}"
                imgs.append({
                    "data_uri": data_uri,
                    "media_type": media_type,
                    "base64": b64_data,      # required by call_claude_for_chunking
                    "para_index": para_idx,  # used for Claude prompt annotation
                })
            except Exception:
                # Fallback: skip sending to Claude but keep in source_html
                imgs.append({
                    "data_uri": src,
                    "media_type": media_type,
                    "base64": None,
                    "para_index": para_idx,
                })
        return imgs

    def is_word_list_para(tag) -> bool:
        classes = tag.get('class') or []
        cls_str = ' '.join(classes) if isinstance(classes, list) else str(classes)
        return 'MsoListParagraph' in cls_str or 'MsoListBullet' in cls_str

    def is_all_bold(tag) -> bool:
        text = tag.get_text().strip()
        if not text:
            return False
        bold_text = ''.join(b.get_text() for b in tag.find_all(['b', 'strong']))
        for span in tag.find_all('span'):
            style = (span.get('style', '') or '').replace(' ', '').lower()
            if 'font-weight:bold' in style or 'font-weight:700' in style:
                bold_text += span.get_text()
        return len(bold_text.strip()) >= len(text) * 0.8

    def process_element(elem):
        nonlocal para_idx
        if not isinstance(elem, Tag):
            return
        tag = elem.name or ''

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            text = elem.get_text(' ', strip=True)
            html_out = node_to_clean_html(elem)
            imgs = extract_images_from_node(elem)
            if text or imgs:
                elements.append({
                    'type': 'heading', 'text': text, 'html': html_out or text,
                    'bold_terms': [text] if text else [], 'level': int(tag[1]),
                    'para_index': para_idx, 'images': imgs, 'is_empty': False,
                })
                all_images.extend(imgs)
                para_idx += 1

        elif tag == 'table':
            rows = []
            for tr in elem.find_all('tr'):
                cells = []
                for td in tr.find_all(['td', 'th']):
                    cells.append({'text': td.get_text(' ', strip=True), 'html': node_to_clean_html(td)})
                if cells:
                    rows.append(cells)
            if rows:
                elements.append({
                    'type': 'table',
                    'text': '\n'.join(' | '.join(c['text'] for c in row) for row in rows),
                    'html': table_to_html(rows), 'bold_terms': [], 'level': 0,
                    'para_index': para_idx, 'images': [], 'is_empty': False,
                })
                para_idx += 1

        elif tag in ('ul', 'ol'):
            # ── Level calculation helper ──────────────────────────────────────
            def get_li_margin(li_elem) -> float:
                # Priority 1: CSS class from <style> block (Pages: li.li1 {margin: ... left})
                classes = li_elem.get('class') or []
                for cls in (classes if isinstance(classes, list) else [classes]):
                    for key in (f"{li_elem.name}.{cls}", f".{cls}"):
                        if key in css_margins:
                            return css_margins[key]
                # Priority 2: inline margin-left / padding-left
                style = li_elem.get('style', '') or ''
                m = re.search(r'(?:margin|padding)-left\s*:\s*([\d.]+)', style, re.I)
                if m:
                    return float(m.group(1))
                return 0.0

            # Build margin→offset map across ALL <li> descendants of this list
            # so relative levels are consistent even with Pages' flat+nested hybrid
            all_margins_here = sorted({get_li_margin(li) for li in elem.find_all('li')})
            margin_to_offset = {v: i for i, v in enumerate(all_margins_here)}

            # Base level: count ancestor ul/ol that are NOT direct-ul-in-ul (Pages artefact)
            proper_nesting = sum(
                1 for a in elem.parents
                if a.name in ('ul', 'ol') and a.parent and a.parent.name not in ('ul', 'ol')
            )
            base_level = 3 + proper_nesting

            def process_li(li_elem):
                nonlocal para_idx
                # Google Docs provides aria-level on <li> — use it as authoritative source
                aria = li_elem.get('aria-level')
                if aria is not None:
                    try:
                        level = int(aria) + 2  # aria-level 1 → our level 3
                    except (ValueError, TypeError):
                        margin = get_li_margin(li_elem)
                        level = base_level + margin_to_offset.get(margin, 0)
                else:
                    margin = get_li_margin(li_elem)
                    level = base_level + margin_to_offset.get(margin, 0)

                direct_text_parts = []
                for child in li_elem.children:
                    if isinstance(child, NavigableString):
                        direct_text_parts.append(str(child))
                    elif isinstance(child, Tag) and child.name not in ('ul', 'ol'):
                        direct_text_parts.append(child.get_text(' '))
                text = ' '.join(direct_text_parts).strip()

                html_parts_li = []
                for child in li_elem.children:
                    if isinstance(child, NavigableString):
                        html_parts_li.append(html_lib.escape(str(child)))
                    elif isinstance(child, Tag) and child.name not in ('ul', 'ol'):
                        html_parts_li.append(node_to_clean_html(child))
                html_out = ''.join(html_parts_li).strip()

                imgs = extract_images_from_node(li_elem)
                if text or imgs:
                    elements.append({
                        'type': 'bullet', 'text': text, 'html': html_out or text,
                        'bold_terms': extract_bold_terms(li_elem), 'level': level,
                        'para_index': para_idx, 'images': imgs, 'is_empty': False,
                    })
                    all_images.extend(imgs)
                    para_idx += 1
                # Recurse into standard nested lists inside this li.
                # Use process_element (not process_ul_in_order) so base_level is
                # recomputed from the actual ancestor depth for each sub-list.
                for nested in li_elem.find_all(['ul', 'ol'], recursive=False):
                    process_element(nested)

            def process_ul_in_order(ul_elem):
                """Process a <ul>/<ol> in DOM order, handling Pages' <ul>-in-<ul> pattern."""
                for child in ul_elem.children:
                    if not isinstance(child, Tag):
                        continue
                    if child.name == 'li':
                        process_li(child)
                    elif child.name in ('ul', 'ol'):
                        # Pages: <ul> directly inside <ul> — recurse in order
                        process_ul_in_order(child)

            process_ul_in_order(elem)

        elif tag == 'p':
            text = elem.get_text(' ', strip=True)
            # Skip blank / non-breaking-space-only Word paragraphs
            if not text or set(text) <= {'\xa0', ' ', '\u200b'}:
                return
            html_out = node_to_clean_html(elem)
            imgs = extract_images_from_node(elem)

            # Detect § bullet prefix (Word/Google Docs list items rendered as <p>)
            is_section_bullet = text.startswith('§')
            if is_section_bullet:
                text = re.sub(r'^§\s*', '', text)
                html_out = re.sub(r'§\s*', '', html_out, count=1)
                # Compute level from margin-left (each indent level ≈ 48px)
                style = elem.get('style', '') or ''
                ml_match = re.search(r'margin-left\s*:\s*([\d.]+)\s*(pt|px)', style, re.I)
                if ml_match:
                    margin_val = float(ml_match.group(1))
                    if ml_match.group(2).lower() == 'pt':
                        margin_val *= 1.333  # pt → px
                    level = max(3, round(margin_val / 48) + 2)
                else:
                    level = 4  # default sub-bullet
                elem_type = 'bullet'
            elif is_word_list_para(elem):
                style = elem.get('style', '') or ''
                level_m = re.search(r'level(\d+)', style)
                level = int(level_m.group(1)) + 2 if level_m else 3
                # Strip Word list prefix characters (·, •, –, etc.) from text/html
                text = re.sub(r'^[\u00b7\u2022\u2013\u2014\-\*]\s*', '', text)
                elem_type = 'bullet'
            elif is_all_bold(elem) and len(text) < 120:
                elem_type = 'heading'
                level = 2
            else:
                elem_type = 'paragraph'
                level = 0

            if text or imgs:
                elements.append({
                    'type': elem_type, 'text': text, 'html': html_out or text,
                    'bold_terms': extract_bold_terms(elem), 'level': level,
                    'para_index': para_idx, 'images': imgs, 'is_empty': False,
                })
                all_images.extend(imgs)
                para_idx += 1

        elif tag in ('div', 'section', 'article', 'main', 'body', 'span',
                     'header', 'footer', 'nav', 'aside', 'figure'):
            for child in elem.children:
                if isinstance(child, Tag):
                    process_element(child)

        elif tag in ('b', 'strong', 'i', 'em'):
            # Google Docs wraps its entire clipboard payload in <b id="docs-internal-guid-...">
            # Detect block-container usage (contains p/ul/div children) and recurse
            has_block = any(
                isinstance(c, Tag) and c.name in ('p', 'ul', 'ol', 'table', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6')
                for c in elem.children
            )
            if has_block:
                for child in elem.children:
                    if isinstance(child, Tag):
                        process_element(child)

    body = soup.find('body')
    root = body if body else soup
    for child in root.children:
        if isinstance(child, Tag):
            process_element(child)

    return elements, all_images


def parse_and_chunk_html(html: str, client: anthropic.Anthropic, model: str = "claude-sonnet-4-6", curriculum_leaves: list[dict] = None) -> tuple[list, dict]:
    """Parse clipboard HTML and chunk it semantically.

    When curriculum_leaves is provided, uses topic-first slicing (slice_by_topics)
    which assigns topics during chunking. Otherwise falls back to AI-driven
    semantic chunking via call_claude_for_chunking.
    """
    elements, images = parse_html_to_elements(html)
    if not elements:
        raise ValueError("No content could be extracted from the pasted HTML")

    if curriculum_leaves:
        chunks, usage = slice_by_topics(elements, curriculum_leaves, client, model=model)
        return chunks, usage

    claude_chunks, usage = call_claude_for_chunking(elements, images, None, client, model=model)
    chunks = assemble_chunks(elements, claude_chunks)
    return chunks, usage


def slice_by_topics(
    elements: list,
    curriculum_leaves: list[dict],  # [{id, name, path, sort_order}, ...]
    client: anthropic.Anthropic,
    model: str = "claude-sonnet-4-6",
) -> tuple[list[dict], dict]:
    """Topic-first slicing: match document headings against curriculum leaves to slice
    the document into topic-aligned chunks. Returns (chunks, usage).

    Algorithm:
    1. Extract all heading-like elements with their indices
    2. Send headings + curriculum tree to Claude in one batch call
    3. Claude returns heading->topic mapping with confidence levels
    4. Use the mapping to slice elements between headings into topic chunks
    """
    # ── 1. Collect headings ──────────────────────────────────────────────
    headings = []
    for i, elem in enumerate(elements):
        if elem["type"] == "heading":
            headings.append({"index": i, "text": elem["text"], "level": elem.get("level", 0)})

    if not headings:
        # No headings found — treat entire document as one chunk, no topic
        chunk = _build_heuristic_chunk(0, 0, len(elements) - 1, "Document", elements)
        chunk["topic_id"] = None
        chunk["topic_path"] = None
        chunk["needs_review"] = True
        return [chunk], {"input_tokens": 0, "output_tokens": 0}

    # ── 2. Build the Claude prompt ───────────────────────────────────────
    curriculum_lines = []
    for leaf in curriculum_leaves:
        curriculum_lines.append(f"  ID={leaf['id']}  path=\"{leaf['path']}\"  name=\"{leaf['name']}\"")
    curriculum_block = "\n".join(curriculum_lines)

    heading_lines = []
    for h in headings:
        # Include a snippet of content after the heading for disambiguation
        after_text = ""
        for j in range(h["index"] + 1, min(h["index"] + 4, len(elements))):
            if elements[j]["type"] != "heading":
                after_text += " " + elements[j]["text"][:100]
        heading_lines.append(
            f"  heading_index={h['index']}  level={h['level']}  "
            f"text=\"{h['text']}\""
            f"  context_after=\"{after_text.strip()[:200]}\""
        )
    heading_block = "\n".join(heading_lines)

    prompt = f"""You are matching document headings from medical study notes to curriculum topics.

IMPORTANT ASSUMPTIONS:
- Topics appear in curriculum order approximately 90% of the time. Use sequential order as a STRONG prior.
- The curriculum leaves below are listed in their natural curriculum order (by sort_order/tree position).
- Each heading should map to exactly one curriculum leaf topic.

DISAMBIGUATION RULES:
- Some topic names appear under different parents (e.g., "Treatment" under both "Pneumonia" and "Asthma").
- When a heading is ambiguous, look at what comes AFTER it (the context_after field) to determine which branch we're in.
- Also use the sequential position: if the previous heading matched topic X, the next heading likely matches a topic near X in curriculum order.

CURRICULUM LEAVES (in curriculum order):
{curriculum_block}

DOCUMENT HEADINGS (in document order):
{heading_block}

For each heading, assign the best matching curriculum leaf.
Return a COMPACT JSON array — only include heading_index, topic_id, and confidence to minimize output size:
[
  [<heading_index>, <topic_id_or_null>, "<h|m|l>"],
  ...
]

Use "h" for high confidence, "m" for medium, "l" for low.

Rules:
- Every heading must appear in the output
- If a heading is a top-level category/title (e.g., "CARDIOVASCULAR") that doesn't map to a specific leaf, use the FIRST leaf id under that category, with "m"
- If a heading truly has no matching topic, use null for topic_id with "l"

Return ONLY the JSON array, no other text."""

    response = client.messages.create(
        model=model,
        max_tokens=16384,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )

    usage = {
        "input_tokens": response.usage.input_tokens,
        "output_tokens": response.usage.output_tokens,
    }

    if response.stop_reason == "max_tokens":
        raise ValueError(
            "Topic slicing response was truncated — document has too many headings for single-pass slicing"
        )

    response_text = response.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = re.sub(r"^```\w*\n?", "", response_text)
        response_text = re.sub(r"\n?```$", "", response_text)

    try:
        heading_mappings = json.loads(response_text)
    except json.JSONDecodeError as e:
        logger.error("slice_by_topics: Claude returned non-JSON: %s", response_text[:300])
        raise ValueError(f"Topic slicing returned non-JSON: {e}") from e

    # ── 3. Build index lookup for fast access ────────────────────────────
    # Build a topic_id → {name, path} lookup from curriculum_leaves
    leaf_lookup = {leaf["id"]: leaf for leaf in curriculum_leaves}
    confidence_map = {"h": "high", "m": "medium", "l": "low"}

    mapping_by_index = {}
    for item in heading_mappings:
        # Compact format: [heading_index, topic_id_or_null, "h"|"m"|"l"]
        if isinstance(item, list):
            h_idx = item[0]
            t_id = item[1]
            conf = confidence_map.get(item[2], "low") if len(item) > 2 else "low"
            leaf = leaf_lookup.get(t_id, {})
            mapping_by_index[h_idx] = {
                "heading_index": h_idx,
                "topic_id": t_id,
                "topic_path": leaf.get("path"),
                "confidence": conf,
            }
        else:
            # Legacy dict format fallback
            mapping_by_index[item["heading_index"]] = item

    # ── 4. Slice elements into topic chunks ──────────────────────────────
    # Each heading starts a new chunk; everything between two headings belongs
    # to the first heading's topic.
    heading_indices = sorted(mapping_by_index.keys())
    chunks = []

    for pos, h_idx in enumerate(heading_indices):
        # Determine the element range for this chunk
        start = h_idx
        if pos + 1 < len(heading_indices):
            end = heading_indices[pos + 1] - 1
        else:
            end = len(elements) - 1

        if end < start:
            continue

        mapping = mapping_by_index[h_idx]
        chunk_elements = elements[start: end + 1]

        # Infer content_type from elements (same logic as heuristic_chunk)
        has_bullets = any(e["type"] == "bullet" for e in chunk_elements)
        has_paragraphs = any(e["type"] == "paragraph" for e in chunk_elements)
        has_table = any(e["type"] == "table" for e in chunk_elements)
        has_images = any(e.get("images") for e in chunk_elements)

        if has_table:
            content_type = "table"
        elif has_images:
            content_type = "reference-image"
        elif has_bullets and has_paragraphs:
            content_type = "mixed-paragraph-bullets"
        elif has_bullets:
            content_type = "bullet-list"
        else:
            content_type = "paragraph"

        topic_id = mapping.get("topic_id")
        topic_path = mapping.get("topic_path")
        confidence = mapping.get("confidence", "low")

        # Use the document heading text as the chunk heading
        heading_text = elements[h_idx]["text"]

        chunk = {
            "chunk_index": len(chunks),
            "heading": heading_text,
            "content_type": content_type,
            "rule_subset": ["cloze_boundaries"],
            "element_range": [start, end],
            "topic_id": topic_id,
            "topic_path": topic_path,
            "needs_review": confidence == "low",
            "topic_confidence": confidence,
        }
        chunks.append(chunk)

    # Handle elements before the first heading (preamble)
    if heading_indices and heading_indices[0] > 0:
        preamble_end = heading_indices[0] - 1
        preamble_elems = elements[0: preamble_end + 1]
        # Skip if it's all empty/whitespace
        if any(e["text"].strip() for e in preamble_elems):
            # Assign preamble to the same topic as the first heading
            first_mapping = mapping_by_index[heading_indices[0]]
            preamble_chunk = {
                "chunk_index": -1,  # will be renumbered below
                "heading": "Introduction",
                "content_type": "paragraph",
                "rule_subset": ["cloze_boundaries"],
                "element_range": [0, preamble_end],
                "topic_id": first_mapping.get("topic_id"),
                "topic_path": first_mapping.get("topic_path"),
                "needs_review": True,
                "topic_confidence": "low",
            }
            chunks.insert(0, preamble_chunk)

    # Renumber chunk indices
    for i, chunk in enumerate(chunks):
        chunk["chunk_index"] = i

    # ── 5. Assemble final chunks with content ────────────────────────────
    result = []
    for cc in chunks:
        start, end = cc["element_range"]
        chunk_elements = elements[start: end + 1]
        source_text_parts = []
        source_html_parts = []
        all_bold_terms = []
        all_images = []

        for elem in chunk_elements:
            source_text_parts.append(elem["text"])
            if elem["type"] == "heading":
                level = min(elem.get("level", 2), 4)
                source_html_parts.append(f"<h{level + 1}>{elem['html']}</h{level + 1}>")
            elif elem["type"] == "table":
                source_html_parts.append(elem["html"])
            elif elem["type"] == "bullet":
                source_html_parts.append(
                    f'<div class="bullet level-{elem.get("level", 3)}">{elem["html"]}</div>'
                )
            else:
                source_html_parts.append(f"<p>{elem['html']}</p>")
            all_bold_terms.extend(elem.get("bold_terms", []))
            for img in elem.get("images", []):
                all_images.append({
                    "data_uri": img["data_uri"],
                    "type": "reference-image",
                    "ocr_text": None,
                    "position": f"in_element_{elem['para_index']}",
                })
                source_html_parts.append(f'<img src="{img["data_uri"]}" class="doc-image" />')

        seen = set()
        unique_bold = []
        for t in all_bold_terms:
            if t not in seen:
                seen.add(t)
                unique_bold.append(t)

        source_html = "\n".join(source_html_parts)
        img_match = IMG_DATA_URI_RE.search(source_html)

        result.append({
            "chunk_index": cc["chunk_index"],
            "heading": cc["heading"],
            "content_type": cc["content_type"],
            "rule_subset": cc["rule_subset"],
            "source_text": "\n".join(source_text_parts),
            "source_html": source_html,
            "bold_terms": unique_bold,
            "element_range": cc["element_range"],
            "images": all_images,
            "ref_img": img_match.group(1) if img_match else None,
            "topic_id": cc.get("topic_id"),
            "topic_path": cc.get("topic_path"),
            "needs_review": cc.get("needs_review", False),
            "topic_confidence": cc.get("topic_confidence", "high"),
        })

    return result, usage


def strip_images_for_storage(chunks: list[dict]) -> list[dict]:
    """Replace inline base64 data URIs in source_html with placeholder text."""
    result = []
    for chunk in chunks:
        new_chunk = dict(chunk)
        new_chunk["source_html"] = re.sub(
            r'<img src="data:[^"]+base64,[^"]*"',
            '<img src="[image]"',
            chunk.get("source_html", "")
        )
        result.append(new_chunk)
    return result


def heuristic_chunk(elements: list) -> list:
    """Fallback heuristic chunking when Claude API is not used."""
    chunks = []
    current_chunk_elements = []
    current_heading = "Introduction"
    chunk_idx = 0
    start_idx = 0

    for i, elem in enumerate(elements):
        if elem["type"] == "heading" and current_chunk_elements:
            chunks.append(_build_heuristic_chunk(
                chunk_idx, start_idx, i - 1, current_heading, current_chunk_elements
            ))
            chunk_idx += 1
            current_chunk_elements = []
            start_idx = i

        if elem["type"] == "heading":
            current_heading = elem["text"]

        current_chunk_elements.append(elem)

    if current_chunk_elements:
        chunks.append(_build_heuristic_chunk(
            chunk_idx, start_idx, len(elements) - 1, current_heading, current_chunk_elements
        ))

    return chunks


def _build_heuristic_chunk(chunk_idx, start, end, heading, elems):
    """Build a chunk dict from elements for heuristic mode."""
    has_bullets = any(e["type"] == "bullet" for e in elems)
    has_paragraphs = any(e["type"] == "paragraph" for e in elems)
    has_table = any(e["type"] == "table" for e in elems)
    has_images = any(e.get("images") for e in elems)

    if has_table:
        content_type = "table"
    elif has_images:
        content_type = "reference-image"
    elif has_bullets and has_paragraphs:
        content_type = "mixed-paragraph-bullets"
    elif has_bullets:
        content_type = "bullet-list"
    else:
        content_type = "paragraph"

    source_text_parts = []
    source_html_parts = []
    all_bold_terms = []
    all_images = []

    for elem in elems:
        source_text_parts.append(elem["text"])
        if elem["type"] == "heading":
            source_html_parts.append(f"<h3>{elem['html']}</h3>")
        elif elem["type"] == "table":
            source_html_parts.append(elem["html"])
        elif elem["type"] == "bullet":
            source_html_parts.append(
                f'<div class="bullet level-{elem.get("level", 3)}">{elem["html"]}</div>'
            )
        else:
            source_html_parts.append(f"<p>{elem['html']}</p>")
        all_bold_terms.extend(elem.get("bold_terms", []))
        for img in elem.get("images", []):
            all_images.append({
                "data_uri": img["data_uri"],
                "type": "reference-image",
                "ocr_text": None,
                "position": f"in_element_{elem['para_index']}",
            })
            source_html_parts.append(f'<img src="{img["data_uri"]}" class="doc-image" />')

    seen = set()
    unique_bold = []
    for t in all_bold_terms:
        if t not in seen:
            seen.add(t)
            unique_bold.append(t)

    return {
        "chunk_index": chunk_idx,
        "heading": heading,
        "content_type": content_type,
        "rule_subset": ["cloze_boundaries"],
        "source_text": "\n".join(source_text_parts),
        "source_html": "\n".join(source_html_parts),
        "bold_terms": unique_bold,
        "element_range": [start, end],
        "images": all_images,
    }


def parse_and_chunk_docx(docx_path: str, img_dir: str, client: anthropic.Anthropic, rules_md_path: Optional[str] = None, model: str = "claude-sonnet-4-6", curriculum_leaves: list[dict] = None) -> tuple[list[dict], dict]:
    """Full pipeline: parse docx -> chunk -> return (assembled chunks, usage).

    When curriculum_leaves is provided, uses topic-first slicing (slice_by_topics)
    which assigns topics during chunking. Otherwise falls back to AI-driven
    semantic chunking via call_claude_for_chunking.
    """
    elements, images = parse_docx(docx_path, img_dir)

    if curriculum_leaves:
        chunks, usage = slice_by_topics(elements, curriculum_leaves, client, model=model)
        return chunks, usage

    claude_chunks, usage = call_claude_for_chunking(elements, images, rules_md_path, client, model=model)
    chunks = assemble_chunks(elements, claude_chunks)
    return chunks, usage
